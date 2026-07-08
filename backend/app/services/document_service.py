"""文档服务 - 文档上传、解析、分块、向量化"""

import hashlib
import os
from typing import Optional

from fastapi import UploadFile
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.exceptions import NotFoundException
from app.models.document import Document, DocumentChunk
from app.models.user import User


def compute_file_hash(content: bytes) -> str:
    """计算文件 SHA-256 哈希"""
    return hashlib.sha256(content).hexdigest()


async def save_uploaded_file(upload_file: UploadFile) -> tuple[str, bytes]:
    """保存上传文件到临时目录，返回文件路径和内容"""
    os.makedirs("./data/uploads", exist_ok=True)
    file_path = f"./data/uploads/{upload_file.filename}"
    content = await upload_file.read()

    with open(file_path, "wb") as f:
        f.write(content)

    return file_path, content


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """从不同类型的文件中提取文本"""
    if file_type == "txt" or file_type == "md":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif file_type == "pdf":
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    elif file_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    elif file_type == "csv":
        import csv
        import io
        with open(file_path, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
            return "\n".join(rows)

    elif file_type == "xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)
        wb.close()
        return "\n".join(text_parts)

    else:
        # 默认按纯文本读取
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()


def clean_text(text: str) -> str:
    """清洗文本：去除多余空白、规范化"""
    import re
    # 去除多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 去除行内多余空格
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # 去除控制字符（保留换行和制表符）
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text.strip()


async def create_document_record(
    db: AsyncSession,
    filename: str,
    file_type: str,
    file_size: int,
    file_hash: str,
    uploader: User,
) -> Document:
    """创建文档数据库记录"""
    collection_name = f"doc_{file_hash[:16]}"
    document = Document(
        filename=filename,
        file_type=file_type,
        file_size=file_size,
        file_hash=file_hash,
        uploader_id=uploader.id,
        collection_name=collection_name,
        status="processing",
    )
    db.add(document)
    await db.flush()
    await db.refresh(document)
    return document


async def get_document(db: AsyncSession, document_id: int) -> Document:
    """获取文档详情"""
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if doc is None:
        raise NotFoundException("文档不存在")
    return doc


async def list_documents(
    db: AsyncSession, page: int = 1, page_size: int = 20
) -> tuple[list[Document], int]:
    """分页获取文档列表"""
    count_query = select(func.count()).select_from(Document)
    total = (await db.execute(count_query)).scalar() or 0

    query = (
        select(Document)
        .order_by(Document.created_at.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
    )
    result = await db.execute(query)
    return list(result.scalars().all()), total


async def delete_document(db: AsyncSession, document_id: int) -> Document:
    """删除文档及其向量数据"""
    doc = await get_document(db, document_id)

    # 删除 ChromaDB 中的向量集合（在 rag_service 中处理）
    # 删除数据库记录（级联删除 chunk）
    await db.delete(doc)
    await db.flush()
    return doc


async def get_knowledge_stats(db: AsyncSession) -> dict:
    """获取知识库统计信息"""
    # 文档总数
    total_docs_result = await db.execute(select(func.count()).select_from(Document))
    total_docs = total_docs_result.scalar() or 0

    # 分块总数
    total_chunks_result = await db.execute(select(func.count()).select_from(DocumentChunk))
    total_chunks = total_chunks_result.scalar() or 0

    # 总字符数
    total_chars_result = await db.execute(select(func.sum(DocumentChunk.char_count)).select_from(DocumentChunk))
    total_chars = total_chars_result.scalar() or 0

    # 按文件类型统计
    type_query = select(Document.file_type, func.count()).group_by(Document.file_type)
    type_result = await db.execute(type_query)
    by_type = {row[0]: row[1] for row in type_result.all()}

    # 按状态统计
    status_query = select(Document.status, func.count()).group_by(Document.status)
    status_result = await db.execute(status_query)
    by_status = {row[0]: row[1] for row in status_result.all()}

    return {
        "total_documents": total_docs,
        "total_chunks": total_chunks,
        "total_chars": total_chars,
        "documents_by_type": by_type,
        "documents_by_status": by_status,
    }
